# -*- coding: utf-8 -*-
"""
Created on Thu May 21 17:27:51 2020

@author: Lenovo
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  #用作设置段落对齐
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
from docx.shared import Inches #图片尺寸

import time
today=time.strftime("%Y-%m-%d",time.localtime())
price=100
company_list=['客户1','客户2','客户3','客户4','客户5']

for i in company_list:
	document =Document()
	document.styles['Normal'].font.name=u'微软雅黑'
	document.styles['Normal'].font.size=Pt(14)
	#设置文档的基础字体
	document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'微软雅黑')
	#添加图片   图片路径和尺寸
	document.add_picture('1.jpg',width=Inches(6))
	p1=document.add_paragraph()
	#初始化建立第一个自然段  设置对齐方式为居中，默认为左对齐
	p1.aligment=WD_ALIGN_PARAGRAPH.CENTER
	#标题内容
	run1=p1.add_run('关于下达%s产品价格的通知'%(today))
	run1.font.name='微软雅黑'
	run1.font.size=Pt(21) #设置字体
	 #设置加粗
	run1.font.bold=True
	#段后距离5磅
	p1.space_after=Pt(5)
	##段前距离5磅
	p1.space_before=Pt(5)
	p2=document.add_paragraph()
	run2=p2.add_run(i+':')
	#这里是对客户的称呼
	run2.font.name='仿宋_GB2312'
	run2.element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋_GB2312')
	run2.font.size=Pt(16) #设置字体
	#设置加粗
	run2.font.bold=True
	p3=document.add_paragraph()
	run3=p3.add_run('   根据公司安排，为提供优质客户服务，我单位拟定了今日黄金价格为%s元，特此通知。'%price)
	#这里是对客户的称呼
	run3.font.name='仿宋_GB2312'
	run2.element.rPr.rFonts.set(qn('w:eastAsia'),u'仿宋_GB2312')
	run3.font.size=Pt(16) #设置字体

	#添加一个表格  行列 和格式
	table=document.add_table(rows=3,cols=3,style='Table Grid')
	#合并单元格
	table.cell(0,0).merge(table.cell(0,2))
	#对于合并的单元格，输入其中任何一个单元格都可以
	table_run1=table.cell(0,0).paragraphs[0].add_run('XX产品报价表')
	table_run1.font.name=u'隶书'
	table_run1.element.rPr.rFonts.set(qn('w:eastAsia'),u'隶书')
	table.cell(0,0).paragraphs[0].aligment=WD_ALIGN_PARAGRAPH.CENTER
	#使用默认字体和格式
	table.cell(1,0).text='日期'
	table.cell(1,1).text='价格'
	table.cell(1,2).text='备注'
	table.cell(2,0).text=today
	table.cell(2,1).text=str(price)
	table.cell(2,2).text=''
	#插入分页符
	document.add_page_break()

	p5=document.add_paragraph()
	run4=p5.add_run('此处是广告')
	document.save('%s-价格通知.docx'%i)#以“客户名-价格通知”作为文件名保存



















